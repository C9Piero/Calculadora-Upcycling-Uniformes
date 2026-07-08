import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

# Configuración de la página web
st.set_page_config(page_title="Calculadora de Sostenibilidad", layout="wide")

# BASE DE DATOS MAESTRA COMPLETA
bd_factores_completa = {
    "Banner": 9.50, "Bata de laboratorio": 6.57, "Bolsas": 8.00, "Camisa": 6.57,
    "Camisa algodón": 5.00, "Camisa drill": 5.90, "Camisa ignífuga": 5.35,
    "Camisa jean / denim": 5.00, "Camisaco": 5.00, "Camisaco drill": 5.90,
    "Camisaco drill con cinta": 5.96, "Casaca": 6.57, "Casaca drill": 5.90,
    "Casaca polar": 6.00, "Casaca polar con cinta reflectiva": 6.05,
    "Casaca térmica": 5.82, "Chaleco": 6.57, "Chaleco con cinta": 6.62,
    "Chaleco de seguridad": 6.62, "Chaleco polar": 6.00, "Chaleco reversible": 6.57,
    "Chompa": 6.57, "Chompa con cinta reflectiva": 6.62, "Chompa Jorge Chavez": 6.30,
    "Chompa Jorge Chavez con cinta reflectiva": 6.30, "Chompa polar": 6.00,
    "Enterizo": 6.57, "Gorro": 7.92, "Impermeable": 9.42, "Mameluco": 6.57,
    "Mameluco acolchado": 5.82, "Mameluco drill": 5.90, "Mameluco jean reflectivo": 5.35,
    "Merma": 6.57, "Overol": 6.57, "Pantalón": 6.57, "Pantalón algodón": 5.00,
    "Pantalón drill": 5.90, "Pantalón drill con cinta": 5.96, "Pantalón ignífugo": 5.35,
    "Pantalón jean": 5.00, "Pantalón jean / drill": 5.45, "Pantalón jean con cinta reflectiva": 5.05,
    "Pantalón polar": 6.00, "Pantalón térmico": 5.82, "Polera": 5.00,
    "Polera polar": 6.00, "Polo": 5.00, "Polo algodón": 5.00, "Polo con cinta reflectiva": 5.05,
    "Polo manga corta": 5.00, "Polo manga larga": 6.80, "Polo manga larga con cinta reflectiva": 6.86,
    "Polo piqué": 5.00, "Short": 5.00, "Toalla": 5.00, "Chaleco Fluorescente": 6.62
}

# Inicializar la lista de prendas en la sesión de la página web
if 'lista_prendas' not in st.session_state:
    st.session_state.lista_prendas = []

st.title("📊 Calculadora de Sostenibilidad Textil")
st.markdown("Ingresa los datos del cliente y añade las prendas recolectadas para calcular el CO2 evitado.")

# --- SECCIÓN 1: DATOS DEL CLIENTE ---
st.subheader("🏢 1. Datos del Cliente")
col_cli1, col_cli2 = st.columns(2)
with col_cli1:
    cliente = st.text_input("Nombre de la Empresa", value="HILTI PERÚ S.A.")
with col_cli2:
    periodo = st.text_input("Periodo / Mes", value="Julio, 2026")

st.markdown("---")

# --- SECCIÓN 2: FORMULARIO Y TABLA ---
col1, col2 = st.columns([1, 2])

with col1:
    st.subheader("👕 2. Cargar Prendas")
    prenda_sel = st.selectbox("Selecciona la Prenda", options=sorted(list(bd_factores_completa.keys())))
    cantidad = st.number_input("Cantidad (unidades)", min_value=1, value=100)
    peso_total = st.number_input("Peso TOTAL del lote (kg)", min_value=0.1, value=43.0)
    
    col_btn1, col_btn2 = st.columns(2)
    with col_btn1:
        if st.button("➕ Añadir Prenda", use_container_width=True):
            factor = bd_factores_completa.get(prenda_sel, 5.0)
            st.session_state.lista_prendas.append({
                "Tipo de prenda": prenda_sel,
                "Cantidad (und)": int(cantidad),
                "Peso total (kg)": round(peso_total, 2),
                "CO2 Evitado (kg)": round(peso_total * factor, 2)
            })
            st.toast(f"✅ {prenda_sel} añadida")
            
    with col_btn2:
        if st.button("🗑️ Limpiar Todo", use_container_width=True):
            st.session_state.lista_prendas = []
            st.toast("Lista vaciada")

with col2:
    st.subheader("📋 3. Tabla del Proyecto Actual")
    if st.session_state.lista_prendas:
        df = pd.DataFrame(st.session_state.lista_prendas)
        st.dataframe(df, use_container_width=True, hide_index=True)
    else:
        st.info("La tabla está vacía. Añade prendas desde el formulario de la izquierda.")

st.markdown("---")

# --- SECCIÓN 3: ENTREGABLES ---
st.subheader("🚀 4. Entregables")
if st.session_state.lista_prendas:
    df_final = pd.DataFrame(st.session_state.lista_prendas)
    total_und = int(df_final["Cantidad (und)"].sum())
    total_kg = round(df_final["Peso total (kg)"].sum(), 2)
    total_co2 = round(df_final["CO2 Evitado (kg)"].sum(), 2)
    
    st.metric(label="CO2 Total Evitado", value=f"{total_co2} kg CO2e")
    
    # Generar el documento en memoria para descarga directa
    doc = Document()
    doc.add_heading("CONSTANCIA DE VALORIZACIÓN TEXTIL", 0)
    doc.add_paragraph(f"Cliente: {cliente}")
    doc.add_paragraph(f"Periodo: {periodo}")
    doc.add_paragraph(f"\nResultados acumulados:")
    doc.add_paragraph(f"• Cantidad de prendas: {total_und} und")
    doc.add_paragraph(f"• Peso textil total: {total_kg} kg")
    doc.add_paragraph(f"• CO2 evitado por upcycling: {total_co2} kg CO2e")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    st.download_button(
        label="📥 DESCARGAR DOCUMENTO WORD",
        data=buffer,
        file_name=f"Constancia_{cliente.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary"
    )
else:
    st.warning("Añade al menos una prenda para habilitar la descarga del reporte Word.")
